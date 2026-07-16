from pathlib import Path

from docx import Document as WordDocument

from .pdf_parser import PageText, extract_pdf


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


def extract_document(path: Path, extension: str) -> list[PageText]:
    extension = extension.lower()
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".docx":
        return extract_docx(path)
    if extension in {".md", ".txt"}:
        return extract_text(path)
    raise ValueError(f"不支持的文件格式：{extension}")


def extract_docx(path: Path) -> list[PageText]:
    document = WordDocument(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            blocks.append(f"\n{text}\n")
        else:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    text = "\n".join(blocks).strip()
    return [PageText(page_number=1, text=text)] if text else []


def extract_text(path: Path) -> list[PageText]:
    raw = path.read_bytes()
    text = None
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise ValueError("无法识别文本文件编码，请使用 UTF-8 编码")
    text = text.strip()
    return [PageText(page_number=1, text=text)] if text else []
