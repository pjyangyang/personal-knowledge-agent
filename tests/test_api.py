from io import BytesIO
import json

import fitz
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.config import settings
from app.db import Base, get_db
from app.main import app


def pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def docx_bytes(text: str) -> bytes:
    stream = BytesIO()
    document = WordDocument()
    document.add_heading("Research Notes", level=1)
    document.add_paragraph(text)
    document.save(stream)
    return stream.getvalue()


def test_health_and_pdf_query(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}", connect_args={"check_same_thread": False})
    testing_session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    Base.metadata.create_all(bind=engine)
    monkeypatch.setattr(settings, "storage_dir", tmp_path / "uploads")
    settings.storage_dir.mkdir()
    monkeypatch.setattr("app.api.vector_store.index_chunks", lambda *args, **kwargs: None)
    monkeypatch.setattr("app.api.vector_store.search", lambda *args, **kwargs: [])
    monkeypatch.setattr("app.api.vector_store.delete_document", lambda *args, **kwargs: None)
    monkeypatch.setattr("app.api.vector_store.delete_knowledge_base", lambda *args, **kwargs: None)
    def fake_fetch_webpage(url, path):
        text = "This webpage explains retrieval augmented generation and source citations."
        path.write_text(text, encoding="utf-8")
        return "RAG Guide", text
    monkeypatch.setattr("app.api.fetch_webpage", fake_fetch_webpage)

    def override_get_db():
        db = testing_session()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_get_db
    try:
        with TestClient(app) as client:
            assert client.get("/health").json()["status"] == "ok"
            kb = client.post("/api/knowledge-bases", json={"name": "论文资料库"})
            assert kb.status_code == 201
            kb_id = kb.json()["id"]
            skills = client.get("/api/skills")
            assert skills.status_code == 200
            assert any(skill["id"] == "paper_analysis" for skill in skills.json())
            renamed = client.patch(f"/api/knowledge-bases/{kb_id}", json={"name": "联邦学习资料库"})
            assert renamed.json()["name"] == "联邦学习资料库"
            response = client.post(
                f"/api/knowledge-bases/{kb_id}/documents",
                files={"file": ("paper.pdf", BytesIO(pdf_bytes("Federated learning protects raw data.")), "application/pdf")},
            )
            assert response.status_code == 201
            text_document = client.post(
                f"/api/knowledge-bases/{kb_id}/documents",
                files={"file": ("notes.txt", BytesIO("RAG combines retrieval with generation.".encode()), "text/plain")},
            )
            assert text_document.status_code == 201
            assert text_document.json()["source_type"] == "txt"
            word_document = client.post(
                f"/api/knowledge-bases/{kb_id}/documents",
                files={"file": ("research.docx", BytesIO(docx_bytes("The experiment uses a private dataset.")), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert word_document.status_code == 201
            assert word_document.json()["source_type"] == "docx"
            webpage = client.post(
                f"/api/knowledge-bases/{kb_id}/webpages",
                json={"url": "https://example.com/rag"},
            )
            assert webpage.status_code == 201
            assert webpage.json()["source_type"] == "webpage"
            assert webpage.json()["source_url"] == "https://example.com/rag"
            query = client.post(f"/api/knowledge-bases/{kb_id}/query", json={"question": "raw data", "skill_id": "paper_analysis"})
            assert query.status_code == 200
            assert query.json()["evidence_found"] is True
            assert query.json()["citations"][0]["page_number"] == 1
            assert query.json()["conversation_id"] > 0
            assert query.json()["skill_id"] == "paper_analysis"
            assert query.json()["evidence_audit"]["total_claims"] >= 1
            assert query.json()["evidence_audit"]["verdict"] in {
                "grounded", "partially_grounded", "ungrounded", "no_evidence"
            }
            conversation_id = query.json()["conversation_id"]
            history = client.get(f"/api/conversations/{conversation_id}")
            assert history.status_code == 200
            assert len(history.json()["messages"]) == 2
            follow_up = client.post(
                f"/api/knowledge-bases/{kb_id}/query",
                json={"question": "raw data", "conversation_id": conversation_id},
            )
            assert follow_up.status_code == 200
            summary = client.post(f"/api/knowledge-bases/{kb_id}/summarize", json={})
            assert summary.status_code == 200
            assert summary.json()["evidence_found"] is True
            with client.stream(
                "POST",
                f"/api/knowledge-bases/{kb_id}/query/stream",
                json={"question": "raw data", "top_k": 3},
            ) as streamed:
                assert streamed.status_code == 200
                events = [json.loads(line) for line in streamed.iter_lines() if line]
            assert events[0]["type"] == "meta"
            assert any(event["type"] == "token" for event in events)
            assert any(event["type"] == "audit" for event in events)
            assert events[-1]["type"] == "done"
    finally:
        app.dependency_overrides.clear()
